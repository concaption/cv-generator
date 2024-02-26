"""
This file contains the utility functions that are used in the main application.
"""
import uuid
import os
import logging

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx2pdf import convert

from google.cloud import storage
from google.oauth2 import service_account

from app.utils import convert_docx_to_pdf


logger = logging.getLogger(__name__)



class ASI_CV:
    """
    This class is used to create a CV for an ASI employee.
    """
    def __init__(self):
        self.doc = Document()
        self.file_id = str(uuid.uuid4())
        self.qualifications = []
        self.technical_skills = []
        self.languages = []
        self.countries = []
        self.summary_of_experience = []
        self.experiences = []
        self.filename = None
        self.docx_file = None
        
    def _add_name_title(self, name, title, create = False):
        self.name = name
        self.title = title
        if create:
            # TODO: Add the name to the document
            pass

    def _add_qualification(self, degree, field, istitution, year, create = False):
        self.qualifications.append({
            "Degree": degree,
            "Field": field,
            "Institution": istitution,
            "Year": year
        })
        if create:
            # TODO: Add the qualification to the document
            pass

    def _add_technical_skill(self, skill, create = False):
        self.technical_skills.append(skill)
        if create:
            # TODO: Add the technical skill to the document
            pass

    def _add_language(self, language, proficiency, create = False):
        self.languages.append({
            "Language": language,
            "Proficiency": proficiency
        })
        if create:
            # TODO: Add the language to the document
            pass

    def _add_country(self, country, create = False):
        self.countries.append(country)
        if create:
            # TODO: Add the country to the document
            pass

    def _add_summary_of_experience(self, summary, create = False):
        self.summary_of_experience.append(summary)
        if create:
            # TODO: Add the summary of experience to the document
            pass

    def _add_experience(self, date_range, position, organisation, location, summary, is_selected, create = False):
        self.experiences.append({
            "Date Range": date_range,
            "Position": position,
            "Organisation": organisation,
            "Location": location,
            "Summary": summary,
            "IsSelected": is_selected
        })
        if create:
            # TODO: Add the experience to the document
            pass

    def generate_cv(self, filename=None, file_format="pdf", output_type="url", save=False, bucket_name=None, folder=None, credentials=None):
        """
        This function is used to generate the CV for the ASI employee.
        """
        if file_format not in ["docx", "pdf"]:
            raise ValueError("The file format should be either 'docx' or 'pdf'")
        if output_type not in ["url", "file"]:
            raise ValueError("The output type should be either 'url' or 'file'")
        self.setup_document()
        self.add_table()
        self.add_heading("Summary of Experience")
        for summary in self.summary_of_experience:
            self.add_paragraph(summary, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        self.add_heading("Employment History")
        self.add_employment_table()
        self.add_heading("Selected Experience")
        selected_experiences = [experience for experience in self.experiences if experience.get("IsSelected") is True]
        for experience in selected_experiences:
            self.add_heading(experience["Position"] + ", " + experience["Organisation"] + ", " + experience["Location"] + " (" + experience["Date Range"] + ")", line=False)
            self.add_paragraph(experience["Summary"], alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        if output_type == "url":
            if bucket_name is None or folder is None or credentials is None:
                raise ValueError("The bucket name, folder and credentials should be provided when the output type is 'url'")        
            credentials = service_account.Credentials.from_service_account_file(credentials)
            google_storage_client = storage.Client(credentials=credentials)
            bucket = google_storage_client.get_bucket(bucket_name)
            if file_format == "docx":
                file_bytes = self.save_docx(filename=filename, save=False)
            if file_format == "pdf":
                file_bytes = self.save_pdf(filename=filename, save=False)
            blob = bucket.blob(folder + "/" + self.file_id + "." + file_format)
            blob.upload_from_string(file_bytes, content_type="application/" + file_format)
            blob.make_public()
            return blob.public_url
        if output_type == "file":
            if folder is None:
                folder = "outputs"
            if file_format == "docx":
                return self.save_docx(filename=filename, save=save, folder=folder)
            if file_format == "pdf":
                return self.save_pdf(filename=filename, save=save, folder=folder)

    def setup_document(self):
        self.set_margins()

        noraml = self.doc.styles['Normal']
        noraml.font.name = 'Arial'
        noraml.font.size = Pt(10)
        noraml.paragraph_format.space_before = Pt(0)
        noraml.paragraph_format.space_after = Pt(0)
        noraml.paragraph_format.line_spacing = 1.08
    
    def set_margins(self, margins=[0.5, 0.5, 0.5, 0.5]):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(margins[0])
            section.bottom_margin = Inches(margins[1])
            section.left_margin = Inches(margins[2])
            section.right_margin = Inches(margins[3])

    def save_docx(self, filename=None, save=False, folder='outputs'):
        if filename is None:
            self.filename = self.file_id + ".docx"
        else:
            self.filename = filename
        # TODO: Output the file to a folder
        self.doc.save(self.filename)
        # covert the docx file into bytes
        with open(self.filename, "rb") as file:
            file_bytes = file.read()
        if not save:
            os.remove(self.filename)
        self.docx_file = file_bytes
        return file_bytes

    def save_pdf(self, filename=None, save=False, folder='outputs'):
        if filename is None:
            self.filename = self.file_id + ".docx"
        else:
            self.filename = filename
        # TODO: Output the file to a folder
        pdf_file = self.filename.replace(".docx", ".pdf")
        # check if the os is windows
        # TODO: Check if the system has MS Word installed
        if os.name == 'nt':
            print("Windows")
            # When using system that has MS Word installed
            from docx2pdf import convert
            convert(self.filename, pdf_file)
        else:
            print("Linux")
            # When using system that does not have MS Word installed
            self.save_docx(self.filename, save=True)
            convert_docx_to_pdf(self.filename, pdf_file)
        with open(pdf_file, "rb") as file:
            file_bytes = file.read()
        if not save:
            os.remove(pdf_file)
        os.remove(self.filename)
        return file_bytes
    
    def add_table(self):
        table = self.doc.add_table(rows=0, cols=4)
        table.style = 'Table Grid'
        # Define the border color and cell background color
        border_color = "000000"  # Black color
        cell_background_color = "FFFFFF"  # White color

        # Loop through each cell and modify its properties
        for row in table.rows:
            for cell in row.cells:
                # Set cell borders
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = tcPr.get_or_add_tcBorders()
                for border in ['top', 'start', 'bottom', 'end']:
                    element = getattr(tcBorders, f'get_or_add_{border}')()
                    element.set('w:val', 'single')  # set border style
                    element.set('w:sz', '4')  # set border size
                    element.set('w:color', border_color)  # set border color
                
                # Set cell background color
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cell_background_color))
                tcPr.append(shading_elm)
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(2)
        table.columns[2].width = Inches(1.5)
        table.columns[3].width = Inches(2.5)
        # add padding to the left and right of the table as table is getting outside of the page
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.add_table_row_with_two_columns(table, ["Name", "Position"], [[self.name], [self.title]], 'D9E2F3', bold=False, bullet=False)
        formated_qualifications = [ qualification.get("Degree") + " in " + qualification.get("Field") + " from " + qualification.get("Institution") + " in " + qualification.get("Year") for qualification in self.qualifications]
        formated_languages = [ language.get("Language") + " (" + language.get("Proficiency") + ")" for language in self.languages]
        self.add_table_row(table, "Qualification", formated_qualifications, 'D9E2F3', bold=False, bullet=True)
        self.add_table_row(table, "Countries", self.countries, 'D9E2F3', bold=False, bullet=False)
        self.add_table_row(table, "Technical Skills", self.technical_skills, 'D9E2F3', bold=False, bullet=False)
        self.add_table_row(table, "Language Skills", formated_languages, 'D9E2F3', bold=False, bullet=False)

    def add_paragraph(self, text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = alignment
        return p

    def add_heading(self, text, line=True):
        heading = self.doc.add_heading(text, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(4)
        heading.paragraph_format.line_spacing = 1.1
        heading.paragraph_format.left_indent = Inches(0)
        run = heading.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if line:
            self.add_horizontal_line(heading)
        return heading
    
    def add_employment_table(self, experiences=None):
        employment_table = self.doc.add_table(rows=0, cols=2)
        employment_table.style = 'TableNormal'
        employment_table.columns[0].width = Inches(1.5)
        employment_table.columns[1].width = Inches(6)
        employment_table.autofit = True
        employment_table.autofit = True
        if experiences is None:
            experiences = self.experiences
        if len(experiences) > 0:
            for experience in experiences:
                row_cells = employment_table.add_row().cells
                self.add_shaded_cell(row_cells[0], experience["Date Range"], "FFFFFF", bold=False)
                self.add_shaded_cell(row_cells[1], f"{experience.get('Position')}, {experience.get('Organisation')}, {experience.get('Location')}", "FFFFFF", bold=False)
        return employment_table

    def add_bullet_point(self, cell, text):
        """
        Adds a bullet point to a cell in a table
        """
        if cell.paragraphs[0].text == '':
            p = cell.paragraphs[0]
            p.clear()
        else:
            p = cell.add_paragraph()
        p.style = 'ListBullet'
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(text)
        return p
    
    def add_shaded_cell(self, cell, text, shade, bold=False):
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.bold = bold
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shade))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        return cell

    def add_table_row(self, table, heading, list, shade, bold=False, bullet=False):
        row_cells = table.add_row().cells
        self.add_shaded_cell(row_cells[0], heading, shade, bold=bold)
        # set the line spacing of all the cells in the row
        # for cell in row_cells:
        #     for paragraph in cell.paragraphs:
        #         paragraph_format = paragraph.paragraph_format
        #         paragraph_format.space_before = Pt(2)
        #         paragraph_format.space_after = Pt(2)
        #         paragraph_format.line_spacing = 1.5
        if not bullet:
            row_cells[1].text = ', '.join(list)
        if bullet:
            texts = [paragraph.text for paragraph in row_cells[1].paragraphs]
            # remove all the paragraphs in the cell
            for item in list:
                self.add_bullet_point(row_cells[1], item)
            texts = [paragraph.text for paragraph in row_cells[1].paragraphs]
        a = row_cells[1]
        b = row_cells[2]
        c = row_cells[3]
        a.merge(b)
        a.merge(c)

    def add_table_row_with_two_columns(self, table, headings, lists, shade, bold=False, bullet=False):
        row_cells = table.add_row().cells
        if len(headings) != 2 or len(lists) != 2:
            raise ValueError("The headings and lists should have two items each")
        self.add_shaded_cell(cell=row_cells[0], text=headings[0], shade=shade, bold=bold)
        self.add_shaded_cell(row_cells[1], lists[0][0],"FFFFFF")
        self.add_shaded_cell(row_cells[2], headings[1], shade, bold=bold)
        self.add_shaded_cell(row_cells[3], lists[1][0],"FFFFFF")

    # unused
    def add_simple_heading(self, text, line=True):
        p = self.doc.add_paragraph(text)
        # add style to the paragraph
        p.style = 'Normal Heading 1'
        if line:
            self.add_horizontal_line(p)
        return p

    def add_horizontal_line(self, p, color="000000", width=4, space=1):
        # add a horizontal line
        p_border = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')  # Single line
        bottom_border.set(qn('w:sz'), str(width))  # Size of the border, e.g., a value of 4 is 1/2 point
        bottom_border.set(qn('w:space'), str(space))  # The space above the border, e.g., 1/8 point
        bottom_border.set(qn('w:color'), color)  # The color of the border
        p_border.append(bottom_border)
        p._p.get_or_add_pPr().append(p_border)